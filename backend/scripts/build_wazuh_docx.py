"""Gera docs/guia-wazuh/GUIA_INSTALACAO_WAZUH.docx a partir do markdown."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs"
MD_PATH = DOCS_DIR / "guia-wazuh" / "GUIA_INSTALACAO_WAZUH.md"
OUT_PATH = DOCS_DIR / "guia-wazuh" / "GUIA_INSTALACAO_WAZUH.docx"

IMAGE_RE = re.compile(r"^!\[(.+?)\]\((.+?)\)\s*$")
FIGURE_WIDTH = Inches(6.2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_image(doc: Document, caption: str, rel_path: str) -> None:
    img_path = (DOCS_DIR / rel_path).resolve()
    if not img_path.is_file():
        p = doc.add_paragraph(f"[Imagem não encontrada: {rel_path}]")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=FIGURE_WIDTH)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(12)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = row[j] if j < len(row) else ""


def add_paragraph_with_formatting(doc: Document, text: str) -> None:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def build_docx(md_text: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        img_match = IMAGE_RE.match(line.strip())
        if img_match and not in_code:
            add_image(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line[1:]:
            if is_table_separator(line):
                i += 1
                continue
            table_rows.append(parse_table_row(line))
            i += 1
            if i < len(lines) and not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue
        elif table_rows:
            add_table(doc, table_rows)
            table_rows = []

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.strip().startswith("- [ ]"):
            text = line.strip()[6:].strip()
            doc.add_paragraph(f"☐ {text}", style="List Bullet")
            i += 1
            continue
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            doc.add_paragraph(m.group(2), style="List Number")
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith(">"):
            p = doc.add_paragraph(line.strip().lstrip("> ").strip())
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        add_paragraph_with_formatting(doc, line.strip())
        i += 1

    if table_rows:
        add_table(doc, table_rows)
    if code_buf:
        add_code_block(doc, code_buf)

    return doc


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    doc = build_docx(md)
    doc.core_properties.title = "Guia Instalação Wazuh — INCIDENT_SYS"
    doc.core_properties.subject = "Wazuh SIEM + INCIDENT_SYS"
    doc.core_properties.author = "INCIDENT_SYS"
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Gerado: {OUT_PATH}")


if __name__ == "__main__":
    main()
